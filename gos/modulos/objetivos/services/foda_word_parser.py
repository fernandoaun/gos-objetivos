"""Extrae ítems FODA desde un documento Word (.docx)."""
import re
from dataclasses import dataclass

from docx import Document
from docx.table import Table

# Palabras clave → tipo (orden: más específico primero)
_SECTION_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r"\bfortalezas?\b", re.I), "F"),
    (re.compile(r"\bstrengths?\b", re.I), "F"),
    (re.compile(r"\boportunidades?\b", re.I), "O"),
    (re.compile(r"\bopportunities?\b", re.I), "O"),
    (re.compile(r"\bdebilidades?\b", re.I), "D"),
    (re.compile(r"\bweaknesses?\b", re.I), "D"),
    (re.compile(r"\bamenazas?\b", re.I), "A"),
    (re.compile(r"\bthreats?\b", re.I), "A"),
]

_BULLET_RE = re.compile(r"^[\s]*(?:[-•·*–—]|\d+[.)]\s+|\w[.)]\s+)", re.U)
_HEADING_STYLES = ("heading", "title", "toc")


@dataclass
class ParsedFodaItem:
    tipo: str
    descripcion: str
    orden: int


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _detect_section(line: str) -> str | None:
    clean = _normalize(line)
    if not clean or len(clean) > 120:
        return None
    # Título de sección: corto o termina en dos puntos
    looks_like_heading = (
        len(clean) < 60
        or clean.endswith(":")
        or clean.isupper()
        or re.match(r"^[IVXLC]+\.", clean)
        or re.match(r"^\d+\.?\s*[A-Za-zÁÉÍÓÚáéíóú]+$", clean)
    )
    if not looks_like_heading and ":" not in clean[:40]:
        return None
    for pattern, tipo in _SECTION_PATTERNS:
        if pattern.search(clean):
            return tipo
    return None


def _is_bullet_line(line: str) -> bool:
    return bool(_BULLET_RE.match(line)) or line.strip().startswith(("-", "•", "·"))


def _strip_bullet(line: str) -> str:
    return _BULLET_RE.sub("", line).strip()


def _parse_paragraphs(doc: Document) -> list[ParsedFodaItem]:
    items: list[ParsedFodaItem] = []
    current_tipo: str | None = None
    orden = 0

    for para in doc.paragraphs:
        raw = para.text
        if not raw or not raw.strip():
            continue
        line = _normalize(raw)
        style = (para.style.name or "").lower()

        section = _detect_section(line)
        if section:
            current_tipo = section
            # Texto en la misma línea después de ":"
            if ":" in line:
                rest = line.split(":", 1)[1].strip()
                if len(rest) > 15:
                    orden += 1
                    items.append(ParsedFodaItem(current_tipo, rest, orden))
            continue

        if any(h in style for h in _HEADING_STYLES):
            section = _detect_section(line)
            if section:
                current_tipo = section
                continue

        if not current_tipo:
            continue

        if _is_bullet_line(raw) or len(line) > 12:
            desc = _strip_bullet(line) if _is_bullet_line(raw) else line
            if len(desc) >= 8:
                orden += 1
                items.append(ParsedFodaItem(current_tipo, desc, orden))

    return items


def _parse_tables(doc: Document) -> list[ParsedFodaItem]:
    """Matriz 2x2 o tabla con columnas F/O/D/A."""
    items: list[ParsedFodaItem] = []
    orden = 0

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [_normalize(c.text) for c in table.rows[0].cells]
        col_map: dict[int, str] = {}

        for idx, h in enumerate(header_cells):
            for pattern, tipo in _SECTION_PATTERNS:
                if pattern.search(h):
                    col_map[idx] = tipo
                    break

        if col_map:
            for row in table.rows[1:]:
                for idx, tipo in col_map.items():
                    if idx < len(row.cells):
                        text = _normalize(row.cells[idx].text)
                        if len(text) >= 8:
                            orden += 1
                            items.append(ParsedFodaItem(tipo, text, orden))
            continue

        # Matriz 2x2 clásica (sin encabezado FODA)
        if len(table.rows) >= 2 and len(table.rows[0].cells) >= 2:
            matrix_map = [
                ("F", 0, 0),
                ("D", 0, 1),
                ("O", 1, 0),
                ("A", 1, 1),
            ]
            if len(table.rows) == 2 and len(table.rows[0].cells) == 2:
                for tipo, r, c in matrix_map:
                    if r < len(table.rows) and c < len(table.rows[r].cells):
                        cell_text = _normalize(table.rows[r].cells[c].text)
                        for part in re.split(r"[\n;]+", cell_text):
                            part = _strip_bullet(_normalize(part))
                            if len(part) >= 8:
                                orden += 1
                                items.append(ParsedFodaItem(tipo, part, orden))

    return items


def _merge_items(paragraph_items: list[ParsedFodaItem], table_items: list[ParsedFodaItem]) -> list[ParsedFodaItem]:
    seen: set[tuple[str, str]] = set()
    merged: list[ParsedFodaItem] = []
    orden = 0
    for item in paragraph_items + table_items:
        key = (item.tipo, item.descripcion.lower()[:200])
        if key in seen:
            continue
        seen.add(key)
        orden += 1
        merged.append(ParsedFodaItem(item.tipo, item.descripcion, orden))
    return merged


def parse_foda_docx(file_path: str) -> list[ParsedFodaItem]:
    doc = Document(file_path)
    paragraph_items = _parse_paragraphs(doc)
    table_items = _parse_tables(doc)
    items = _merge_items(paragraph_items, table_items)

    if not items:
        raise ValueError(
            "No se detectaron ítems FODA. Usá títulos como Fortalezas, Oportunidades, "
            "Debilidades y Amenazas, con viñetas o párrafos debajo de cada sección."
        )
    return items
