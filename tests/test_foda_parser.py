import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.services.foda_word_parser import parse_foda_docx


def _make_docx(path: Path):
    doc = Document()
    doc.add_heading("Análisis FODA", 0)
    doc.add_heading("Fortalezas", level=1)
    doc.add_paragraph("Equipo técnico certificado en operaciones.", style="List Bullet")
    doc.add_paragraph("Buena relación con clientes estratégicos.", style="List Bullet")
    doc.add_heading("Oportunidades", level=1)
    doc.add_paragraph("Crecimiento del mercado regional de energía.", style="List Bullet")
    doc.add_heading("Debilidades", level=1)
    doc.add_paragraph("Dependencia de un solo proveedor crítico.", style="List Bullet")
    doc.add_heading("Amenazas", level=1)
    doc.add_paragraph("Mayor regulación ambiental en el sector.", style="List Bullet")
    doc.save(path)


def test_parse_foda_docx():
    with tempfile.TemporaryDirectory() as tmp:
        p = Path(tmp) / "foda.docx"
        _make_docx(p)
        items = parse_foda_docx(str(p))
        tipos = {i.tipo for i in items}
        assert tipos == {"F", "O", "D", "A"}
        assert len(items) >= 5
